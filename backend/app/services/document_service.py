"""Document processing service for RAG."""
import pdfplumber
import docx
import uuid
from pathlib import Path
from sqlalchemy.orm import Session
from typing import Optional

from app.models.knowledge import KnowledgeDocument, DocumentStatus
from app.models.file import File
from app.utils.chunking import chunk_text
from app.services.chroma_service import chroma_service


async def extract_text(file_path: str, content_type: str) -> str:
    """Extract text from different file types."""
    if content_type == "application/pdf":
        return await _extract_pdf(file_path)
    elif content_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
        return await _extract_docx(file_path)
    elif content_type == "text/plain":
        return await _extract_text(file_path)
    else:
        raise ValueError(f"Unsupported file type: {content_type}")


async def _extract_pdf(file_path: str) -> str:
    """Extract text from PDF."""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
                text += "\n"
    except Exception as e:
        raise ValueError(f"Failed to extract PDF: {str(e)}")
    return text


async def _extract_docx(file_path: str) -> str:
    """Extract text from DOCX."""
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        raise ValueError(f"Failed to extract DOCX: {str(e)}")
    return text


async def _extract_text(file_path: str) -> str:
    """Extract text from plain text file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        raise ValueError(f"Failed to extract text: {str(e)}")


async def process_document(
    document_id: str,
    file_path: str,
    file_content_type: str,
    collection_name: str,
    chunk_size: int = 512,
    overlap: int = 50,
    db: Session = None
) -> None:
    """Process a document: extract text, chunk, embed, and upsert to Chroma."""
    doc = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == document_id).first()
    if not doc:
        return

    try:
        # Update status to processing
        doc.status = DocumentStatus.PROCESSING
        db.commit()

        # Extract text
        text = await extract_text(file_path, file_content_type)

        if not text.strip():
            raise ValueError("No text extracted from document")

        # Chunk text
        chunks = chunk_text(text, chunk_size=chunk_size, overlap=overlap)

        # Create embeddings (mock for now - would call embedding API)
        embeddings = []
        for chunk in chunks:
            # In production, call embedding API
            # For now, use mock embeddings
            embedding = [0.0] * 1536  # Mock embedding
            embeddings.append(embedding)

        # Upsert to Chroma
        doc_ids = [f"{document_id}_{i}" for i in range(len(chunks))]
        metadatas = [
            {"doc_id": document_id, "chunk": i, "title": doc.title}
            for i in range(len(chunks))
        ]

        await chroma_service.upsert(
            collection_name=collection_name,
            documents=chunks,
            embeddings=embeddings,
            ids=doc_ids,
            metadatas=metadatas
        )

        # Update document status
        doc.status = DocumentStatus.INDEXED
        doc.chunk_count = len(chunks)
        from datetime import datetime
        doc.indexed_at = datetime.utcnow()
        db.commit()

    except Exception as e:
        # Update error status
        doc.status = DocumentStatus.FAILED
        doc.error_message = str(e)
        db.commit()
